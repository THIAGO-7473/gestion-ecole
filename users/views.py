from pyexpat.errors import messages
from django.contrib import messages
from django.shortcuts import redirect, render
from django.views.generic import ListView, CreateView, UpdateView, DeleteView, View
from django.urls import reverse_lazy
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.db.models import Q
from django.http import HttpResponse, HttpResponseRedirect
from .models import Utilisateur
from .forms import UtilisateurForm
import csv
# from django.template.loader import get_template
# from xhtml2pdf import pisa
import openpyxl
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.template.loader import render_to_string
from widget_tweaks.templatetags.widget_tweaks import add_class
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.views import LoginView, LogoutView
from django.views.generic import TemplateView
from django.utils import timezone
from docx import Document
from docx.shared import Inches
import pdfkit
import os
import pandas as pd
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from io import BytesIO
from reportlab.lib.pagesizes import landscape, letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

class UtilisateurListView(LoginRequiredMixin, ListView):
    model = Utilisateur
    template_name = 'users/utilisateur_list.html'
    context_object_name = 'utilisateurs'
    paginate_by = 10
    ordering = ['-date_creation']

    def get_queryset(self):
        queryset = super().get_queryset()
        search = self.request.GET.get('search')
        role = self.request.GET.get('role')
        statut = self.request.GET.get('statut')

        if search:
            queryset = queryset.filter(
                Q(username__icontains=search) |
                Q(first_name__icontains=search) |
                Q(last_name__icontains=search) |
                Q(email__icontains=search)
            )
        if role and role != 'all':
            queryset = queryset.filter(role=role)
        if statut and statut != 'all':
            queryset = queryset.filter(statut=statut)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['role_choices'] = [('all', 'Tous les rôles')] + list(Utilisateur.Role.choices)
        context['statut_choices'] = [('all', 'Tous les statuts')] + list(Utilisateur.Statut.choices)
        context['current_role'] = self.request.GET.get('role', 'all')
        context['current_statut'] = self.request.GET.get('statut', 'all')
        context['current_search'] = self.request.GET.get('search', '')
        return context

class UtilisateurCreateView(LoginRequiredMixin, UserPassesTestMixin, CreateView):
    model = Utilisateur
    form_class = UtilisateurForm
    template_name = 'users/utilisateur_form.html'
    success_url = reverse_lazy('utilisateur_list')

    def test_func(self):
        return self.request.user.is_admin()

class UtilisateurUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Utilisateur
    form_class = UtilisateurForm
    template_name = 'users/utilisateur_form.html'
    success_url = reverse_lazy('utilisateur_list')

    def test_func(self):
        return self.request.user.is_admin() or self.request.user == self.get_object()

    def form_valid(self, form):
        try:
            # Sauvegarder l'utilisateur
            self.object = form.save()
            # Ajouter un message de succès
            messages.success(self.request, f"L'utilisateur {self.object.username} a été mis à jour avec succès.")
            return HttpResponseRedirect(self.get_success_url())
        except Exception as e:
            messages.error(self.request, f"Une erreur est survenue : {str(e)}")
            return self.form_invalid(form)

    def form_invalid(self, form):
        messages.error(self.request, "Veuillez corriger les erreurs ci-dessous.")
        return super().form_invalid(form)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['instance'] = self.get_object()
        return kwargs

class UtilisateurDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Utilisateur
    template_name = 'users/utilisateur_confirm_delete.html'
    success_url = reverse_lazy('utilisateur_list')
    
    # Implémentation requise pour UserPassesTestMixin
    def test_func(self):
        """Seuls les admins ou l'utilisateur lui-même peuvent supprimer"""
        user_to_delete = self.get_object()
        return self.request.user.is_admin() or self.request.user == user_to_delete
    
    def handle_no_permission(self):
        """Gestion personnalisée des refus d'accès"""
        messages.error(self.request, "Vous n'avez pas la permission de supprimer cet utilisateur.")
        return redirect(self.success_url)
    
    def delete(self, request, *args, **kwargs):
        """Surcharge pour ajouter des messages"""
        messages.success(request, f"L'utilisateur {self.get_object().username} a été supprimé avec succès.")
        return super().delete(request, *args, **kwargs)
    

    
def export_users_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="utilisateurs.csv"'

    writer = csv.writer(response)
    writer.writerow(['N°', 'Username', 'Nom complet', 'Email', 'Rôle', 'Statut', 'Date création'])

    for idx, user in enumerate(Utilisateur.objects.all().order_by('-date_creation'), start=1):
        writer.writerow([
            idx,
            user.username,
            user.nom_complet,
            user.email,
            user.get_role_display(),
            user.get_statut_display(),
            user.date_creation.strftime("%Y-%m-%d")
        ])

    return response

def export_users_pdf(request):
    try:
        # Création du document PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
        elements = []

        # Style pour l'en-tête
        styles = getSampleStyleSheet()
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Centré
        )
        header = Paragraph("Liste des Utilisateurs", header_style)
        elements.append(header)

        # Date de génération
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=20,
            alignment=1  # Centré
        )
        date_text = Paragraph(f"Généré le {timezone.now().strftime('%d/%m/%Y %H:%M')}", date_style)
        elements.append(date_text)

        # Préparation des données pour le tableau
        headers = ['N°', 'Username', 'Nom complet', 'Email', 'Rôle', 'Statut', 'Date création']
        data = [headers]  # En-tête

        for idx, user in enumerate(Utilisateur.objects.all().order_by('-date_creation'), start=1):
            row = [
                str(idx),
                user.username,
                user.nom_complet,
                user.email,
                user.get_role_display(),
                user.get_statut_display(),
                user.date_creation.strftime("%d/%m/%Y")
            ]
            data.append(row)

        # Création du tableau
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]))

        elements.append(table)
        doc.build(elements)
        pdf = buffer.getvalue()
        buffer.close()

        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="utilisateurs.pdf"'
        return response

    except Exception as e:
        messages.error(request, f"Erreur lors de la génération du PDF : {str(e)}")
        return redirect('utilisateur_list')

def export_users_word(request):
    users = Utilisateur.objects.all().order_by('-date_creation')
    
    # Création du document Word
    doc = Document()
    
    # Titre
    doc.add_heading('Liste des Utilisateurs', 0)
    doc.add_paragraph(f'Généré le {timezone.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph()
    
    # Tableau
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # En-têtes
    headers = ['Photo', 'Nom', 'Postnom', 'Prénom', 'Email', 'Rôle', 'Statut']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    
    # Données
    for user in users:
        row = table.add_row()
        # Photo (si disponible)
        if user.photo_profil and os.path.exists(user.photo_profil.path):
            row.cells[0].text = "Photo disponible"
        else:
            row.cells[0].text = "-"
        
        # Autres informations
        row.cells[1].text = user.last_name or "-"
        row.cells[2].text = user.postnom or "-"
        row.cells[3].text = user.first_name or "-"
        row.cells[4].text = user.email or "-"
        row.cells[5].text = user.get_role_display()
        row.cells[6].text = user.get_statut_display()
    
    # Sauvegarde du document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="utilisateurs.docx"'
    doc.save(response)
    
    return response

def export_users_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="utilisateurs.xlsx"'

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Utilisateurs"

    rows = [
        ['N°', 'Username', 'Nom complet', 'Email', 'Rôle', 'Statut', 'Date création']
    ]

    for idx, user in enumerate(Utilisateur.objects.all().order_by('-date_creation'), start=1):
        rows.append([
            idx,
            user.username,
            user.nom_complet,
            user.email,
            user.get_role_display(),
            user.get_statut_display(),
            user.date_creation.strftime("%Y-%m-%d")
        ])

    for row in rows:
        ws.append(row)

    wb.save(response)
    return response

class AccueilView(TemplateView):
    template_name = 'users/accueil.html'

    def get(self, request, *args, **kwargs):
        if request.user.is_authenticated:
            return redirect('utilisateur_list')
        return super().get(request, *args, **kwargs)

class ImportUsersView(LoginRequiredMixin, View):
    template_name = 'users/import_users.html'
    
    def get(self, request):
        context = {
            'role_choices': [choice[0] for choice in Utilisateur.Role.choices],
            'statut_choices': [choice[0] for choice in Utilisateur.Statut.choices],
        }
        return render(request, self.template_name, context)
    
    def post(self, request):
        if 'file' not in request.FILES:
            messages.error(request, 'Aucun fichier n\'a été sélectionné.')
            return redirect('import_users')
        
        file = request.FILES['file']
        file_extension = os.path.splitext(file.name)[1].lower()
        
        try:
            if file_extension == '.csv':
                df = pd.read_csv(file)
            elif file_extension in ['.xlsx', '.xls']:
                df = pd.read_excel(file)
            else:
                messages.error(request, 'Format de fichier non supporté. Utilisez CSV ou Excel.')
                return redirect('import_users')
            
            # Vérifier les colonnes requises
            required_columns = ['last_name', 'postnom', 'first_name', 'email', 'role', 'statut']
            missing_columns = [col for col in required_columns if col not in df.columns]
            
            if missing_columns:
                messages.error(request, f'Colonnes manquantes dans le fichier : {", ".join(missing_columns)}')
                return redirect('import_users')
            
            # Vérifier les valeurs de rôle et statut
            valid_roles = [choice[0] for choice in Utilisateur.Role.choices]
            valid_statuts = [choice[0] for choice in Utilisateur.Statut.choices]
            
            invalid_roles = df[~df['role'].isin(valid_roles)]['role'].unique()
            invalid_statuts = df[~df['statut'].isin(valid_statuts)]['statut'].unique()
            
            if len(invalid_roles) > 0:
                messages.error(request, f'Rôles invalides trouvés : {", ".join(invalid_roles)}')
                return redirect('import_users')
            
            if len(invalid_statuts) > 0:
                messages.error(request, f'Statuts invalides trouvés : {", ".join(invalid_statuts)}')
                return redirect('import_users')
            
            # Traiter chaque ligne
            success_count = 0
            error_count = 0
            errors = []
            
            for index, row in df.iterrows():
                try:
                    # Générer un nom d'utilisateur unique basé sur l'email
                    base_username = row['email'].split('@')[0]
                    username = base_username
                    counter = 1
                    
                    while Utilisateur.objects.filter(username=username).exists():
                        username = f"{base_username}{counter}"
                        counter += 1
                    
                    if Utilisateur.objects.filter(email=row['email']).exists():
                        errors.append(f'Ligne {index + 2}: L\'email {row["email"]} est déjà utilisé.')
                        error_count += 1
                        continue
                    
                    # Créer l'utilisateur
                    user = Utilisateur.objects.create_user(
                        username=username,
                        email=row['email'],
                        password='password123',  # Mot de passe par défaut
                        first_name=row['first_name'],
                        last_name=row['last_name'],
                        postnom=row['postnom'],
                        role=row['role'],
                        statut=row['statut']
                    )
                    
                    success_count += 1
                    
                except Exception as e:
                    errors.append(f'Ligne {index + 2}: {str(e)}')
                    error_count += 1
            
            # Afficher les messages de succès et d'erreur
            if success_count > 0:
                messages.success(request, f'{success_count} utilisateur(s) importé(s) avec succès.')
            
            if errors:
                for error in errors:
                    messages.warning(request, error)
            
            if error_count > 0:
                messages.warning(request, f'{error_count} erreur(s) lors de l\'importation.')
            
        except Exception as e:
            messages.error(request, f'Erreur lors du traitement du fichier: {str(e)}')
        
        return redirect('utilisateur_list')

class UserDashboardView(LoginRequiredMixin, TemplateView):
    template_name = 'users/dashboard.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        user = self.request.user
        
        # Statistiques générales
        context['total_users'] = Utilisateur.objects.count()
        context['active_users'] = Utilisateur.objects.filter(statut='actif').count()
        context['suspended_users'] = Utilisateur.objects.filter(statut='suspendu').count()
        
        # Répartition par rôle
        context['role_distribution'] = {
            'admin': Utilisateur.objects.filter(role='admin').count(),
            'enseignant': Utilisateur.objects.filter(role='enseignant').count(),
            'personnel': Utilisateur.objects.filter(role='personnel').count(),
            'parent': Utilisateur.objects.filter(role='parent').count(),
            'eleve': Utilisateur.objects.filter(role='eleve').count()
        }
        
        # Dernières activités
        context['recent_users'] = Utilisateur.objects.order_by('-date_creation')[:5]
        
        # Informations personnelles
        context['user_info'] = {
            'last_login': user.dernier_login,
            'account_age': (timezone.now() - user.date_creation).days,
            'is_active': user.is_active,
            'role': user.get_role_display(),
            'status': user.get_statut_display()
        }
        
        return context
